#!/usr/bin/env python3
"""
Split the master Glossa Ordinaria docx into individual book files
(mirroring the original site's per-book downloads). Rex pacificus is
included with Book 1, following Prof. Reno's own Book 1 isolate.

Usage: python3 split_books.py MASTER.docx OUTPUT_DIR [suffix]
"""
import sys, os, copy
from docx import Document

GROUPS = [
    ("Book 1", ["REX PACIFICUS", "BOOK I"]),
    ("Book 2", ["BOOK II"]),
    ("Book 3", ["BOOK III"]),
    ("Book 4", ["BOOK IV"]),
    ("Book 5", ["BOOK V"]),
]

def split(master, outdir, suffix="rev. 9.23"):
    os.makedirs(outdir, exist_ok=True)
    src = Document(master)
    # map each body <w:p> element to its book (by H1 boundaries)
    book_of = {}
    cur = None
    for p in src.paragraphs:
        if p.style.name == "Heading 1":
            cur = p.text.strip().upper()
        book_of[id(p._p)] = cur

    for label, h1s in GROUPS:
        doc = Document(master)  # fresh copy each time
        body = doc.element.body
        cur = None
        for p in list(doc.paragraphs):
            if p.style.name == "Heading 1":
                cur = p.text.strip().upper()
            if cur not in h1s:
                body.remove(p._p)
        out = os.path.join(outdir, f"Glossa Ordinaria to the Decretals, {label}, {suffix}.docx")
        doc.save(out)
        kept = Document(out)
        h = [p.text for p in kept.paragraphs if p.style.name == "Heading 1"]
        n4 = sum(1 for p in kept.paragraphs if p.style.name == "Heading 4" and p.text.strip())
        chars = sum(len(p.text.strip()) for p in kept.paragraphs if p.style.name == "Normal" and p.text.strip())
        print(f"  {label}: H1={h} lemmata={n4} chars={chars} -> {os.path.basename(out)}")

if __name__ == "__main__":
    split(sys.argv[1], sys.argv[2], sys.argv[3] if len(sys.argv) > 3 else "rev. 9.23")
